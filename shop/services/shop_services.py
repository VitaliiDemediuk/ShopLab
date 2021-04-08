import openpyxl
import docx
from docx.shared import RGBColor
from shop.models import *
from ShopLabWork.settings import MEDIA_ROOT
from django.db.models import Count, Max

def get_sections_with_categories():
    sections = list(Section.objects.all().values('id', 'name', 'name_for_link'))
    for section in sections:
        section['category'] = list(Category.objects.filter(fk_section_id=section['id']).values('name', 'name_for_link'))
    return sections


def get_brands():
    brands = list(Brand.objects.all().values('name', 'name_for_link', 'photo'))
    return brands


def get_sizes_by_id(id: int):
    sizes = []
    sizes_query_set = SizesGoods.objects.filter(fk_goods_id=id)
    for size in sizes_query_set:
        sizes.append({'id': size.fk_size_id.id, 'name': size.fk_size_id.name})
    return sizes


def get_photos_by_id(id: int):
    photos = list(PhotoForGoods.objects.filter(fk_goods_id=id).values_list('file_name', flat=True))
    return photos


def get_photos_path_by_id(id: int):
    photos = list(PhotoForGoods.objects.filter(fk_goods_id=id).values_list('file_name', flat=True))
    return photos


def get_characteristics_by_id(id: int):
    characteristics = []
    characteristics_query_set = CharacteristicsGoods.objects.filter(fk_goods_id=id)
    for characteristic in characteristics_query_set:
        characteristics.append({'name': characteristic.fk_characteristic_id.name, 'value': characteristic.value})
    return characteristics


def get_goods_by_id(id: int):
    goods_query_set = Goods.objects.filter(pk=id)
    goods = None
    if len(goods_query_set) != 0:
        goods = list(goods_query_set.values('id', 'title', 'price', 'sale_price', 'main_photo',
                                            'in_stock', 'is_enable', 'description'))[0]
        goods['brand'] = str(goods_query_set[0].fk_brand_id)
        goods['category'] = str(goods_query_set[0].fk_category_id)
        goods['color'] = str(goods_query_set[0].fk_color_id)
        goods['sizes'] = get_sizes_by_id(id)
        goods['photos'] = get_photos_by_id(id)
        goods['characteristics'] = get_characteristics_by_id(id)

    return goods


def get_number_of_products_in_each_category():
    sql_query = 'SELECT "shop_goods"."fk_category_id_id" AS "id", ' \
                'CONCAT("shop_section"."name", \' / \' ,"shop_category"."name") AS "full_name", ' \
                'COUNT("shop_goods"."id") AS "number" FROM "shop_goods" ' \
                'INNER JOIN "shop_category" ON ("shop_goods"."fk_category_id_id" = "shop_category"."id") ' \
                'INNER JOIN "shop_section" ON ("shop_category"."fk_section_id_id" = "shop_section"."id") ' \
                'GROUP BY "shop_goods"."fk_category_id_id", "shop_section"."name", "shop_category"."name";'
    number_of_products = Goods.objects.raw(sql_query)
    return number_of_products


def get_number_of_products_in_each_section():
    sql_query = 'SELECT "shop_section"."id" AS "id", ' \
                '"shop_section"."name", ' \
                'COUNT("shop_goods"."id") AS "number" FROM "shop_goods" ' \
                'INNER JOIN "shop_category" ON ("shop_goods"."fk_category_id_id" = "shop_category"."id") ' \
                'INNER JOIN "shop_section" ON ("shop_category"."fk_section_id_id" = "shop_section"."id") ' \
                'GROUP BY "shop_section"."id", "shop_section"."name";'
    number_of_products = Goods.objects.raw(sql_query)
    return number_of_products

# GOODS LIST
def filter_goods_by_section_link_name(goods_list_query_set, section_link_name):
    if section_link_name is not None:
        section_id = Section.objects.filter(name_for_link=section_link_name).values('id')
        if len(section_id) == 0:
            return None
        else:
            category_ids = list(Category.objects.filter(fk_section_id=section_id[0]['id']).values_list('id', flat=True))
            goods_list_query_set = goods_list_query_set.filter(fk_category_id__in=category_ids)
    return goods_list_query_set


def filter_goods_by_category_link_name(goods_list_query_set, category_link_name):
    if category_link_name is not None:
        category_id = list(Category.objects.filter(name_for_link=category_link_name).values_list('id', flat=True))
        if len(category_id) == 0:
            return None
        else:
            goods_list_query_set = goods_list_query_set.filter(fk_category_id__in=category_id)
    return goods_list_query_set


def filter_goods_by_brand_link_name(goods_list_query_set, brand_link_name):
    if brand_link_name is not None:
        brand_id = list(Brand.objects.filter(name_for_link=brand_link_name).values('id'))
        if len(brand_id) == 0:
            return None
        else:
            brand_id = brand_id[0]['id']
            goods_list_query_set = goods_list_query_set.filter(fk_brand_id=brand_id)
    return goods_list_query_set

def get_goods_query_set(section_link_name = None, category_link_name = None, brand_link_name = None):
    goods_list_query_set = Goods.objects.all()
    # Filter by section:
    goods_list_query_set = filter_goods_by_section_link_name(goods_list_query_set, section_link_name)
    if goods_list_query_set is None:
        return None
    # Filter by category:
    goods_list_query_set = filter_goods_by_category_link_name(goods_list_query_set, category_link_name)
    if goods_list_query_set is None:
        return None
    # Filter by brand:
    goods_list_query_set = filter_goods_by_brand_link_name(goods_list_query_set, brand_link_name)
    if goods_list_query_set is None:
        return None

    return goods_list_query_set


def get_goods_list(section_link_name = None, category_link_name = None, brand_link_name = None):
    goods_list_query_set = get_goods_query_set(section_link_name, category_link_name, brand_link_name)
    goods_list = list(goods_list_query_set.filter(is_enable=True).values('id', 'title', 'price',
                                                                         'sale_price', 'main_photo', 'in_stock'))
    return goods_list

# Export file
def get_max_number_characteristics_in_product():
    max = CharacteristicsGoods.objects.values('fk_goods_id')\
                               .annotate(total=Count('fk_characteristic_id'))\
                               .aggregate(max=Max('total'))['max']
    return max if max else 0


def get_max_number_photos_in_product():
    max = PhotoForGoods.objects.values('fk_goods_id')\
                       .annotate(total=Count('file_name'))\
                       .aggregate(max=Max('total'))['max']
    return max if max else 0


def add_hearer_to_spreadsheet(ws_products, column_list, max_number_characteristics, max_number_photos):
    ws_products.append(column_list)
    column_list_number = len(column_list)

    for i in range(max_number_characteristics):
        ws_products.cell(row=1, column=column_list_number + 2 * i + 1,
                         value=f"ch_name_{i + 1}")
        ws_products.cell(row=1, column=column_list_number + 2 * i + 2,
                         value=f"ch_value_{i + 1}")

    for i in range(max_number_photos):
        ws_products.cell(row=1,
                         column=column_list_number + 2 * max_number_characteristics + i + 1,
                         value=f"photo_{i + 1}")


def add_products_to_spreadsheet(ws_products, products, column_list_number, max_number_characteristics):
    for k, product in enumerate(products, start=2):
        product_data_list = [product.id, product.title, product.price, product.sale_price,
                             product.description, product.count, product.sale_price != 0, product.in_stock,
                             product.is_enable, product.fk_brand_id.name, product.fk_category_id.fk_section_id.name,
                             product.fk_category_id.name, product.fk_color_id.name, product.main_photo.url]
        ws_products.append(product_data_list)
        characteristics = get_characteristics_by_id(product_data_list[0])
        for i, characteristic in enumerate(characteristics):
            ws_products.cell(row=k, column=column_list_number + 2 * i + 1,
                             value=f"{characteristic['name']}")
            ws_products.cell(row=k, column=column_list_number + 2 * i + 2,
                             value=f"{characteristic['value']}")

        photos = get_photos_by_id(product_data_list[0])
        for i, photo in enumerate(photos):
            ws_products.cell(row=k,
                             column=column_list_number + 2 * max_number_characteristics + i + 1,
                             value=f"{photo}")


def get_products_workbook(section_link_name = None, category_link_name = None, brand_link_name = None):
    wb = openpyxl.Workbook()
    ws_products = wb.active
    ws_products.title = "Products"

    column_list = ['id', 'title', 'price', 'sale_price', 'description', 'count', 'is_sale', 'in_stock',
                   'is_enable', 'brand', 'section', 'category', 'color', 'main photo']
    max_number_characteristics = get_max_number_characteristics_in_product()
    max_number_photos = get_max_number_photos_in_product()
    add_hearer_to_spreadsheet(ws_products, column_list, max_number_characteristics, max_number_photos)

    products = get_goods_query_set(section_link_name, category_link_name, brand_link_name)
    add_products_to_spreadsheet(ws_products, products, len(column_list), max_number_characteristics)

    return wb


def get_products_document(section_link_name = None, category_link_name = None, brand_link_name = None):
    document = docx.Document()
    products = get_goods_query_set(section_link_name, category_link_name, brand_link_name)
    for product in products:
        if product.is_enable:
            document.add_heading(product.title, 0)
            document.add_picture(f'{product.main_photo.path}', width=docx.shared.Cm(10))
            document.add_heading(f'Price: {product.price}', 2)
            if product.sale_price != 0:
                run = document.add_paragraph().add_run(f'Sale price: {product.sale_price}')
                font = run.font
                font = run.font
                font.color.rgb = RGBColor.from_string('FF0000')
            document.add_heading(f'Color: {product.fk_color_id.name}', 2)
            document.add_heading(f'Brand: {product.fk_brand_id.name}', 2)
            document.add_heading(f'Section: {product.fk_category_id.fk_section_id.name}', 2)
            document.add_heading(f'Category: {product.fk_category_id.name}', 2)
            document.add_heading(f'Description:', 2)
            document.add_paragraph(product.description)
            photos = get_photos_by_id(product.id)
            for photo in photos:
                document.add_picture(f'{MEDIA_ROOT}/{photo}', width=docx.shared.Cm(7))
            characteristics = get_characteristics_by_id(product.id)
            table = document.add_table(rows=len(characteristics), cols=2)
            for i, characteristic in enumerate(characteristics):
                cell = table.cell(i, 0)
                cell.text = characteristic['name']
                cell = table.cell(i, 1)
                cell.text = characteristic['value']

    return document